"""DOCX template processor for generating reports.

Template markers format: {{field_name}}
Supported markers:
    - {{date}} - Current date (auto-generated)
    - {{stock_code}} - Stock code
    - {{stock_name}} - Stock name
    - {{open}} - Opening price
    - {{high}} - Highest price
    - {{low}} - Lowest price
    - {{close}} - Closing price
    - {{volume}} - Trading volume
    - {{amount}} - Trading amount
    - {{turnover_rate}} - Turnover rate
    - {{change_pct}} - Change percentage
    - {{change_amount}} - Change amount
    - {{sma_5}} - 5-day simple moving average
    - {{sma_10}} - 10-day SMA
    - {{sma_20}} - 20-day SMA
    - {{rsi}} - RSI indicator
    - {{macd}} - MACD indicator
    - {{macd_signal}} - MACD signal
    - {{bb_upper}} - Bollinger upper band
    - {{bb_lower}} - Bollinger lower band
"""

import re
from datetime import date, datetime
from pathlib import Path
from typing import Optional, Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..excel.models import StockDataBundle
from ..lstm import calculate_sma, calculate_rsi, calculate_macd, calculate_bollinger_bands


class DocxTemplateProcessor:
    """Process DOCX templates with stock data."""

    # Regex pattern to find template markers
    MARKER_PATTERN = re.compile(r"\{\{(\w+)\}\}")

    def __init__(self):
        """Initialize template processor."""
        pass

    def process(
        self,
        template_path: str,
        bundle: StockDataBundle,
        output_path: str,
        stock_name: Optional[str] = None,
    ) -> None:
        """Process template and generate report.

        Args:
            template_path: Path to the DOCX template file
            bundle: Stock data bundle
            output_path: Path to save the generated report
            stock_name: Optional stock name override
        """
        doc = Document(template_path)
        context = self._build_context(bundle, stock_name)

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, context)

        # Process all tables
        for table in doc.tables:
            self._process_table(table, context)

        # Save the document
        doc.save(output_path)

    def _build_context(
        self, bundle: StockDataBundle, stock_name: Optional[str] = None
    ) -> dict[str, Any]:
        """Build context dictionary from stock data.

        Args:
            bundle: Stock data bundle
            stock_name: Optional stock name

        Returns:
            Dictionary with all template variables
        """
        closes = [d.close for d in bundle.data]
        volumes = [d.volume for d in bundle.data]
        turnover_rates = [d.turnover_rate for d in bundle.data]

        latest = bundle.data[-1]

        # Calculate technical indicators
        sma_5_list = calculate_sma(closes, 5)
        sma_10_list = calculate_sma(closes, 10)
        sma_20_list = calculate_sma(closes, 20)

        rsi_list = calculate_rsi(closes, 14)
        macd_dif_list, macd_dea_list, macd_hist_list = calculate_macd(closes)
        bb_upper_list, bb_middle_list, bb_lower_list = calculate_bollinger_bands(closes, 20)

        def safe_get(lst, idx, default=0.0):
            """Safely get list value."""
            if 0 <= idx < len(lst):
                val = lst[idx]
                return val if val is not None else default
            return default

        latest_idx = len(closes) - 1

        context = {
            # Basic info
            "date": datetime.now().strftime("%Y年%m月%d日"),
            "today": date.today().strftime("%Y-%m-%d"),
            "stock_code": bundle.stock_code,
            "stock_name": stock_name or bundle.stock_name or bundle.stock_code,
            # Latest price data
            "open": f"{latest.open:.2f}",
            "high": f"{latest.high:.2f}",
            "low": f"{latest.low:.2f}",
            "close": f"{latest.close:.2f}",
            "volume": f"{latest.volume:,}",
            "amount": f"{latest.amount:,.2f}",
            "turnover_rate": f"{latest.turnover_rate:.2f}",
            "change_pct": f"{latest.change_pct:.2f}",
            "change_amount": f"{latest.change_amount:.2f}",
            # Technical indicators
            "sma_5": f"{safe_get(sma_5_list, latest_idx):.2f}",
            "sma_10": f"{safe_get(sma_10_list, latest_idx):.2f}",
            "sma_20": f"{safe_get(sma_20_list, latest_idx):.2f}",
            "rsi": f"{safe_get(rsi_list, latest_idx):.2f}",
            "macd": f"{safe_get(macd_dif_list, latest_idx):.4f}",
            "macd_signal": f"{safe_get(macd_dea_list, latest_idx):.4f}",
            "macd_hist": f"{safe_get(macd_hist_list, latest_idx):.4f}",
            "bb_upper": f"{safe_get(bb_upper_list, latest_idx):.2f}",
            "bb_middle": f"{safe_get(bb_middle_list, latest_idx):.2f}",
            "bb_lower": f"{safe_get(bb_lower_list, latest_idx):.2f}",
            # Raw values for calculations
            "_sma_5_raw": safe_get(sma_5_list, latest_idx),
            "_sma_10_raw": safe_get(sma_10_list, latest_idx),
            "_sma_20_raw": safe_get(sma_20_list, latest_idx),
            "_rsi_raw": safe_get(rsi_list, latest_idx),
            "_close_raw": latest.close,
        }
        return context

    def _process_paragraph(self, paragraph: Paragraph, context: dict) -> None:
        """Process a single paragraph, replacing markers.

        Args:
            paragraph: Paragraph to process
            context: Context dictionary with values
        """
        if not paragraph.text:
            return

        # Find all markers in the paragraph
        full_text = paragraph.text
        matches = list(self.MARKER_PATTERN.finditer(full_text))

        if not matches:
            return

        # Build new text with replacements
        new_runs = []
        last_end = 0

        for match in matches:
            # Add text before the match
            if match.start() > last_end:
                new_runs.append((full_text[last_end : match.start()], False))

            marker_name = match.group(1)
            value = context.get(marker_name, match.group(0))
            new_runs.append((str(value), True))
            last_end = match.end()

        # Add remaining text
        if last_end < len(full_text):
            new_runs.append((full_text[last_end:], False))

        # Replace runs
        if new_runs:
            # Keep the first run's formatting, update text
            for i, (text, is_marker) in enumerate(new_runs):
                if i == 0:
                    paragraph.runs[0].text = text
                elif is_marker:
                    # Add a new run with same formatting as first run
                    run = paragraph.add_run(text)
                    run.style = paragraph.runs[0].style if paragraph.runs else None
                else:
                    paragraph.add_run(text)

    def _process_table(self, table: Table, context: dict) -> None:
        """Process a table, replacing markers in cells.

        Args:
            table: Table to process
            context: Context dictionary with values
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph, context)
