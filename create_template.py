"""Create example template files."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


def create_docx_template(output_path: str):
    """Create an example DOCX template document."""
    doc = Document()

    # Title
    title = doc.add_heading('股票分析报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('报告日期: {{date}}')
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Section 1: Basic Information
    doc.add_heading('一、基本信息', level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    info_data = [
        ('股票代码', '{{stock_code}}'),
        ('股票名称', '{{stock_name}}'),
        ('报告日期', '{{today}}'),
    ]

    for i, (label, value) in enumerate(info_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Section 2: Price Data
    doc.add_heading('二、交易数据', level=1)

    price_table = doc.add_table(rows=6, cols=2)
    price_table.style = 'Table Grid'

    price_data = [
        ('开盘价', '{{open}} 元'),
        ('最高价', '{{high}} 元'),
        ('最低价', '{{low}} 元'),
        ('收盘价', '{{close}} 元'),
        ('成交量', '{{volume}} 股'),
        ('成交额', '{{amount}} 元'),
    ]

    for i, (label, value) in enumerate(price_data):
        price_table.rows[i].cells[0].text = label
        price_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Section 3: Indicators
    doc.add_heading('三、技术指标', level=1)

    indicator_table = doc.add_table(rows=5, cols=2)
    indicator_table.style = 'Table Grid'

    indicator_data = [
        ('换手率', '{{turnover_rate}} %'),
        ('涨跌幅', '{{change_pct}} %'),
        ('MA5', '{{sma_5}} 元'),
        ('MA10', '{{sma_10}} 元'),
        ('MA20', '{{sma_20}} 元'),
    ]

    for i, (label, value) in enumerate(indicator_data):
        indicator_table.rows[i].cells[0].text = label
        indicator_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Section 4: Analysis
    doc.add_heading('四、分析结论', level=1)

    analysis = doc.add_paragraph()
    analysis.add_run(
        f'根据对{{stock_name}}（股票代码：{{stock_code}}）的技术分析，'
        f'当前收盘价为 {{close}} 元，换手率为 {{turnover_rate}}%。'
    )

    doc.add_paragraph()
    doc.add_paragraph('注：本报告由程序自动生成。')

    doc.save(output_path)
    print(f'DOCX模板已保存到: {output_path}')


if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, 'templates', '股票分析报告模板.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    create_docx_template(output_path)
